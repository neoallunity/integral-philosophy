#!/usr/bin/env python3
"""
Content integrity validation for Integral Philosophy publishing system.
Ensures consistency between different output formats.
"""

import hashlib
import difflib
from pathlib import Path
from typing import Dict, List, Any, Optional, Set, Tuple
from dataclasses import dataclass
from bs4 import BeautifulSoup
import logging

from .validators import BaseValidator, ValidationResult, ValidationError

# Optional dependencies for extended format support
try:
    import ebooklib
    from ebooklib import epub

    EPUB_AVAILABLE = True
except ImportError:
    EPUB_AVAILABLE = False

try:
    from docx import Document

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)


@dataclass
class ContentChunk:
    """Represents a chunk of content with metadata."""

    text: str
    hash: str
    source_format: str
    location: str  # file path or identifier
    metadata: Dict[str, Any]


class ContentIntegrityValidator(BaseValidator):
    """Validates content integrity across different output formats."""

    def __init__(self, config: Optional[Dict[str, Any]] = None):
        super().__init__(config)
        self.tolerance = self.config.get("tolerance", 0.95)  # 95% similarity threshold

    def validate(self, file_path: Path) -> ValidationResult:
        """Validate content integrity by comparing with reference format."""
        # This method is for compatibility - actual integrity validation
        # happens in validate_integrity_across_formats
        return self._create_result(True, [], {})

    def validate_integrity_across_formats(
        self, formats: Dict[str, Path]
    ) -> ValidationResult:
        """
        Validate content integrity across multiple formats.

        Args:
            formats: Dictionary mapping format names to file paths
                    e.g., {'html': Path('output.html'), 'pdf': Path('output.pdf')}
        """
        errors = []
        stats = {
            "formats_checked": len(formats),
            "content_chunks": 0,
            "similarity_scores": {},
        }

        if len(formats) < 2:
            errors.append(
                ValidationError(
                    severity="warning",
                    message="Need at least 2 formats for integrity validation",
                    file_path="",
                    rule_id="integrity-insufficient-formats",
                )
            )
            return self._create_result(True, errors, stats)

        try:
            # Extract text content from each format
            content_by_format = {}
            for format_name, file_path in formats.items():
                try:
                    content = self._extract_text_content(file_path, format_name)
                    content_by_format[format_name] = content
                    stats["content_chunks"] += len(content)
                except Exception as e:
                    errors.append(
                        ValidationError(
                            severity="error",
                            message=f"Failed to extract content from {format_name}: {str(e)}",
                            file_path=str(file_path),
                            rule_id="integrity-extraction-error",
                        )
                    )

            # Compare content between formats
            format_names = list(content_by_format.keys())
            for i in range(len(format_names)):
                for j in range(i + 1, len(format_names)):
                    format1, format2 = format_names[i], format_names[j]
                    similarity = self._compare_content(
                        content_by_format[format1], content_by_format[format2]
                    )

                    stats["similarity_scores"][f"{format1}_vs_{format2}"] = similarity

                    if similarity < self.tolerance:
                        errors.append(
                            ValidationError(
                                severity="warning",
                                message=f"Content similarity between {format1} and {format2}: {similarity:.2%} (threshold: {self.tolerance:.2%})",
                                file_path="",
                                rule_id="integrity-low-similarity",
                            )
                        )

            # Check for structural consistency
            structural_errors = self._validate_structural_consistency(content_by_format)
            errors.extend(structural_errors)

        except Exception as e:
            errors.append(
                ValidationError(
                    severity="error",
                    message=f"Content integrity validation failed: {str(e)}",
                    file_path="",
                    rule_id="integrity-validation-error",
                )
            )

        return self._create_result(
            is_valid=len([e for e in errors if e.severity == "error"]) == 0,
            errors=errors,
            stats=stats,
        )

    def _extract_text_content(
        self, file_path: Path, format_type: str
    ) -> List[ContentChunk]:
        """Extract text content from various formats."""
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        if format_type.lower() == "html":
            return self._extract_from_html(file_path)
        elif format_type.lower() == "epub":
            if not EPUB_AVAILABLE:
                raise ImportError(
                    "ebooklib is required for EPUB support. Install with: pip install ebooklib"
                )
            return self._extract_from_epub(file_path)
        elif format_type.lower() in ["docx", "doc"]:
            if not DOCX_AVAILABLE:
                raise ImportError(
                    "python-docx is required for DOCX support. Install with: pip install python-docx"
                )
            return self._extract_from_docx(file_path)
        elif format_type.lower() == "tex":
            return self._extract_from_latex(file_path)
        elif format_type.lower() in ["txt", "md"]:
            return self._extract_from_text(file_path)
        else:
            raise ValueError(f"Unsupported format: {format_type}")

    def _extract_from_html(self, file_path: Path) -> List[ContentChunk]:
        """Extract content from HTML file."""
        chunks = []

        with open(file_path, "r", encoding="utf-8") as f:
            soup = BeautifulSoup(f.read(), "html.parser")

        # Extract semantic content
        for tag in soup.find_all(
            ["h1", "h2", "h3", "h4", "h5", "h6", "p", "div", "article", "section"]
        ):
            text = tag.get_text(strip=True)
            if text and len(text) > 10:  # Skip very short text
                chunk = ContentChunk(
                    text=text,
                    hash=hashlib.md5(text.encode()).hexdigest(),
                    source_format="html",
                    location=f"{tag.name}:{tag.get('id', 'no-id')}",
                    metadata={"tag": tag.name, "attributes": dict(tag.attrs)},
                )
                chunks.append(chunk)

        return chunks

    def _extract_from_epub(self, file_path: Path) -> List[ContentChunk]:
        """Extract content from EPUB file."""
        chunks = []

        if not EPUB_AVAILABLE:
            raise ImportError(
                "ebooklib is required for EPUB support. Install with: pip install ebooklib"
            )

        try:
            book = epub.read_epub(str(file_path))

            for item in book.get_items():
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    # Parse XHTML content
                    soup = BeautifulSoup(item.get_content(), "html.parser")

                    # Extract text from similar tags as HTML
                    for tag in soup.find_all(
                        ["h1", "h2", "h3", "h4", "h5", "h6", "p", "div"]
                    ):
                        text = tag.get_text(strip=True)
                        if text and len(text) > 10:
                            chunk = ContentChunk(
                                text=text,
                                hash=hashlib.md5(text.encode()).hexdigest(),
                                source_format="epub",
                                location=f"{item.get_name()}:{tag.name}",
                                metadata={"tag": tag.name, "file": item.get_name()},
                            )
                            chunks.append(chunk)

        except Exception as e:
            logger.error(f"Failed to extract from EPUB {file_path}: {e}")
            raise

        return chunks

    def _extract_from_docx(self, file_path: Path) -> List[ContentChunk]:
        """Extract content from DOCX file."""
        chunks = []

        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for DOCX support. Install with: pip install python-docx"
            )

        try:
            doc = Document(str(file_path))

            for para in doc.paragraphs:
                text = para.text.strip()
                if text and len(text) > 10:
                    chunk = ContentChunk(
                        text=text,
                        hash=hashlib.md5(text.encode()).hexdigest(),
                        source_format="docx",
                        location=f"paragraph:{len(chunks)}",
                        metadata={"style": para.style.name if para.style else "Normal"},
                    )
                    chunks.append(chunk)

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text and len(text) > 10:
                            chunk = ContentChunk(
                                text=text,
                                hash=hashlib.md5(text.encode()).hexdigest(),
                                source_format="docx",
                                location=f"table_cell:{len(chunks)}",
                                metadata={"type": "table_cell"},
                            )
                            chunks.append(chunk)

        except Exception as e:
            logger.error(f"Failed to extract from DOCX {file_path}: {e}")
            raise

        return chunks

    def _extract_from_latex(self, file_path: Path) -> List[ContentChunk]:
        """Extract content from LaTeX file."""
        chunks = []

        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()

        # Remove LaTeX commands but keep text
        import re

        # Remove comments
        content = re.sub(r"%.*$", "", content, flags=re.MULTILINE)

        # Extract content from sections, subsections, paragraphs
        section_pattern = (
            r"\\(section|subsection|subsubsection|paragraph)\*?\{([^}]+)\}"
        )

        for match in re.finditer(section_pattern, content):
            title = match.group(2).strip()
            if title:
                chunk = ContentChunk(
                    text=title,
                    hash=hashlib.md5(title.encode()).hexdigest(),
                    source_format="latex",
                    location=f"{match.group(1)}:{match.start()}",
                    metadata={"command": match.group(1)},
                )
                chunks.append(chunk)

        # Extract paragraph content (simplified)
        paragraphs = re.split(r"\n\s*\n", content)
        for i, para in enumerate(paragraphs):
            # Remove LaTeX commands
            clean_text = re.sub(r"\\[a-zA-Z]+\{([^}]*)\}", r"\1", para)
            clean_text = re.sub(r"[{}]", "", clean_text)
            clean_text = clean_text.strip()

            if clean_text and len(clean_text) > 10 and not clean_text.startswith("\\"):
                chunk = ContentChunk(
                    text=clean_text,
                    hash=hashlib.md5(clean_text.encode()).hexdigest(),
                    source_format="latex",
                    location=f"paragraph:{i}",
                    metadata={"type": "paragraph"},
                )
                chunks.append(chunk)

        return chunks

    def _extract_from_text(self, file_path: Path) -> List[ContentChunk]:
        """Extract content from plain text file."""
        chunks = []

        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()

        # Split by paragraphs (double newlines)
        paragraphs = content.split("\n\n")

        for i, para in enumerate(paragraphs):
            text = para.strip()
            if text and len(text) > 10:
                chunk = ContentChunk(
                    text=text,
                    hash=hashlib.md5(text.encode()).hexdigest(),
                    source_format="text",
                    location=f"paragraph:{i}",
                    metadata={"type": "paragraph"},
                )
                chunks.append(chunk)

        return chunks

    def _compare_content(
        self, content1: List[ContentChunk], content2: List[ContentChunk]
    ) -> float:
        """Compare content similarity between two format extractions."""
        if not content1 or not content2:
            return 0.0

        # Extract text hashes for comparison
        hashes1 = {chunk.hash: chunk.text for chunk in content1}
        hashes2 = {chunk.hash: chunk.text for chunk in content2}

        # Find common content
        common_hashes = set(hashes1.keys()) & set(hashes2.keys())

        # Calculate similarity based on common content
        total_chunks = len(set(hashes1.keys()) | set(hashes2.keys()))

        if total_chunks == 0:
            return 0.0

        exact_match_score = len(common_hashes) / total_chunks

        # For non-exact matches, check text similarity
        similar_chunks = 0
        for chunk1 in content1:
            if chunk1.hash not in hashes2:
                # Try to find similar text in content2
                for chunk2 in content2:
                    similarity = difflib.SequenceMatcher(
                        None, chunk1.text, chunk2.text
                    ).ratio()
                    if similarity > 0.8:  # 80% text similarity threshold
                        similar_chunks += 0.5  # Partial credit
                        break

        final_score = exact_match_score + (similar_chunks / total_chunks)
        return min(final_score, 1.0)  # Cap at 1.0

    def _validate_structural_consistency(
        self, content_by_format: Dict[str, List[ContentChunk]]
    ) -> List[ValidationError]:
        """Validate structural consistency across formats."""
        errors = []

        # Check heading hierarchy consistency
        for format_name, content in content_by_format.items():
            headings = [
                chunk
                for chunk in content
                if chunk.metadata.get("tag") in ["h1", "h2", "h3", "h4", "h5", "h6"]
            ]

            # Check for proper heading order (no skipping levels)
            prev_level = 0
            for heading in headings:
                tag = heading.metadata.get("tag", "")
                if tag.startswith("h"):
                    level = int(tag[1:])
                    if level > prev_level + 1:
                        errors.append(
                            ValidationError(
                                severity="warning",
                                message=f"Heading level skip in {format_name}: from h{prev_level} to {tag}",
                                file_path=heading.location,
                                rule_id="integrity-heading-skip",
                            )
                        )
                    prev_level = level

        # Check for missing important content chunks
        all_hashes = set()
        hash_sources = {}

        for format_name, content in content_by_format.items():
            for chunk in content:
                all_hashes.add(chunk.hash)
                if chunk.hash not in hash_sources:
                    hash_sources[chunk.hash] = []
                hash_sources[chunk.hash].append(format_name)

        # Find content that exists in some formats but not others
        missing_content = {
            h: sources
            for h, sources in hash_sources.items()
            if len(sources) < len(content_by_format)
        }

        if missing_content:
            errors.append(
                ValidationError(
                    severity="info",
                    message=f"Found {len(missing_content)} content chunks not present in all formats",
                    file_path="",
                    rule_id="integrity-missing-content",
                )
            )

        return errors


class CrossReferenceValidator(BaseValidator):
    """Validates cross-references between formats."""

    def validate(self, file_path: Path) -> ValidationResult:
        """Validate cross-references - requires multiple files."""
        return self._create_result(True, [], {})

    def validate_cross_references(self, formats: Dict[str, Path]) -> ValidationResult:
        """Validate cross-references between formats."""
        errors = []
        stats = {
            "formats_checked": len(formats),
            "references_found": 0,
            "broken_references": 0,
        }

        # This would need to be implemented based on the specific
        # cross-referencing system used in the publishing pipeline

        return self._create_result(True, errors, stats)
